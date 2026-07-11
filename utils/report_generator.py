from io import BytesIO

from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION


def heading(doc, text, level=1):

    h = doc.add_heading(level=level)

    run = h.add_run(text)

    run.bold = True

    run.font.color.rgb = RGBColor(0, 70, 140)


def normal(doc, text):

    p = doc.add_paragraph()

    r = p.add_run(text)

    r.font.size = Pt(11)


def metric(doc, title, value):

    p = doc.add_paragraph()

    p.style = "List Bullet"

    r = p.add_run(f"{title}: ")

    r.bold = True

    p.add_run(str(value))

def add_cover_page(doc, company):

    title = doc.add_heading(
        "Enterprise AI Transformation Report",
        level=0
    )

    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p = doc.add_paragraph()

    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p.add_run(company["Company"]).bold = True

    p.add_run("\n")

    p.add_run(company["Industry"])

    p.add_run("\n")

    p.add_run("Confidential")

    doc.add_page_break()

def add_executive_summary(doc, report, roi, knowledge):

    heading(doc, "1. Executive Summary")

    normal(
        doc,
        "This assessment evaluates the organisation's AI readiness, "
        "business opportunities, financial impact and recommended "
        "implementation roadmap."
    )

    table = doc.add_table(rows=6, cols=2)

    table.style = "Table Grid"

    table.cell(0,0).text = "Overall AI Potential"
    table.cell(0,1).text = str(report["overall_ai_potential"])

    table.cell(1,0).text = "Estimated Investment"
    table.cell(1,1).text = f"${roi['investment']:,.0f}"

    table.cell(2,0).text = "Annual Benefit"
    table.cell(2,1).text = f"${roi['benefit']:,.0f}"

    table.cell(3,0).text = "ROI"
    table.cell(3,1).text = f"{roi['roi']:.1f}%"

    table.cell(4,0).text = "Payback Period"
    table.cell(4,1).text = f"{roi['payback']:.2f} Years"

    table.cell(5,0).text = "AI Readiness"
    table.cell(5,1).text = f"{knowledge['overall_score']}%"

    heading(doc, "Executive Recommendation", level=2)

    if roi["roi"] >= 100:

        normal(
            doc,
            "The assessment indicates a strong business case for AI "
            "investment. The recommended approach is to begin with "
            "high-priority initiatives in Phase 1 while establishing "
            "governance and an AI Centre of Excellence."
        )

    else:

        normal(
            doc,
            "The organisation should improve data readiness and "
            "business processes before making large AI investments."
        )

    doc.add_page_break()

def add_company_profile(doc, company):

    heading(doc, "2. Company Profile")

    profile = company["company_profile"]

    table = doc.add_table(rows=0, cols=2)

    table.style = "Table Grid"

    for key, value in profile.items():

        row = table.add_row().cells

        row[0].text = str(key)

        row[1].text = str(value)

    doc.add_page_break()

def add_opportunities(doc, report):

    heading(doc, "3. AI Opportunity Assessment")

    normal(
        doc,
        "The following opportunities have been identified as the "
        "highest-value AI initiatives for the organisation."
    )

    table = doc.add_table(rows=1, cols=6)

    table.style = "Table Grid"

    hdr = table.rows[0].cells

    hdr[0].text = "Department"
    hdr[1].text = "Use Case"
    hdr[2].text = "Priority"
    hdr[3].text = "Impact"
    hdr[4].text = "ROI"
    hdr[5].text = "Phase"

    for item in report["top_opportunities"]:

        row = table.add_row().cells

        row[0].text = item["department"]

        row[1].text = item["use_case"]

        row[2].text = item["priority"]

        row[3].text = str(item["business_impact"])

        row[4].text = item["estimated_roi"]

        row[5].text = item["implementation_phase"]

    doc.add_page_break()

def add_roi(doc, roi):

    heading(doc, "4. Financial Business Case")

    table = doc.add_table(rows=6, cols=2)

    table.style = "Table Grid"

    table.cell(0,0).text = "Investment"
    table.cell(0,1).text = f"${roi['investment']:,.0f}"

    table.cell(1,0).text = "Annual Benefit"
    table.cell(1,1).text = f"${roi['benefit']:,.0f}"

    table.cell(2,0).text = "ROI"
    table.cell(2,1).text = f"{roi['roi']:.1f}%"

    table.cell(3,0).text = "Payback"
    table.cell(3,1).text = f"{roi['payback']:.2f} Years"

    table.cell(4,0).text = "Hours Saved"
    table.cell(4,1).text = f"{roi['hours_saved']:,.0f}"

    table.cell(5,0).text = "Salary Savings"
    table.cell(5,1).text = f"${roi['salary_savings']:,.0f}"

    heading(doc, "Financial Interpretation", level=2)

    if roi["roi"] > 200:

        normal(
            doc,
            "The proposed investment delivers an exceptional financial "
            "return with rapid payback and should be prioritised."
        )

    elif roi["roi"] > 100:

        normal(
            doc,
            "The investment demonstrates a strong financial case and "
            "is recommended for implementation."
        )

    else:

        normal(
            doc,
            "Further validation of assumptions is recommended before "
            "approving the investment."
        )

    doc.add_page_break()

def add_roadmap(doc, roadmap):

    heading(doc, "5. AI Transformation Roadmap")

    normal(doc, roadmap["executive_summary"])

    phases = [
        ("Phase 1", roadmap["phase1"]),
        ("Phase 2", roadmap["phase2"]),
        ("Phase 3", roadmap["phase3"])
    ]

    for phase_name, initiatives in phases:

        heading(doc, phase_name, level=2)

        if not initiatives:
            normal(doc, "No initiatives identified.")
            continue

        for initiative in initiatives:

            p = doc.add_paragraph(style="List Bullet")
            p.add_run(str(initiative))

    doc.add_page_break()

def add_readiness(doc, knowledge):

    heading(doc, "6. AI Readiness Assessment")

    metric(
        doc,
        "Overall AI Readiness Score",
        f"{knowledge['overall_score']}%"
    )

    metric(
        doc,
        "Maturity Level",
        knowledge["maturity"]
    )

    heading(doc, "Capability Assessment", level=2)

    table = doc.add_table(rows=1, cols=2)

    table.style = "Table Grid"

    hdr = table.rows[0].cells

    hdr[0].text = "Capability"

    hdr[1].text = "Score"

    for capability, score in knowledge["scores"].items():

        row = table.add_row().cells

        row[0].text = capability

        row[1].text = f"{score}%"

    doc.add_page_break()

def add_portfolio(doc, report, roi):

    heading(doc, "7. AI Investment Portfolio")

    opportunities = report["top_opportunities"]

    investment = roi["investment"] / max(len(opportunities), 1)

    benefit = roi["benefit"] / max(len(opportunities), 1)

    table = doc.add_table(rows=1, cols=6)

    table.style = "Table Grid"

    hdr = table.rows[0].cells

    hdr[0].text = "Initiative"

    hdr[1].text = "Department"

    hdr[2].text = "Investment"

    hdr[3].text = "Benefit"

    hdr[4].text = "Priority"

    hdr[5].text = "Phase"

    for item in opportunities:

        row = table.add_row().cells

        row[0].text = item["use_case"]

        row[1].text = item["department"]

        row[2].text = f"${investment:,.0f}"

        row[3].text = f"${benefit:,.0f}"

        row[4].text = item["priority"]

        row[5].text = item["implementation_phase"]

    doc.add_page_break()

def add_recommendations(doc, report, roi, knowledge):

    heading(doc, "8. Executive Recommendations")

    recommendations = []

    if roi["roi"] >= 100:
        recommendations.append(
            "Approve Phase-1 AI initiatives immediately as the projected ROI is strong."
        )
    else:
        recommendations.append(
            "Revalidate financial assumptions before approving investment."
        )

    if knowledge["overall_score"] < 70:
        recommendations.append(
            "Improve AI readiness through governance, training and data quality initiatives."
        )

    recommendations.append(
        "Establish an AI Centre of Excellence (CoE) to govern enterprise AI initiatives."
    )

    recommendations.append(
        "Measure business value quarterly using defined KPIs."
    )

    recommendations.append(
        "Prioritize high-impact, low-effort use cases before expanding to enterprise-wide AI."
    )

    for rec in recommendations:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(rec)

def generate_report(data):

    doc = Document()

    company = data["company_data"]
    report = data["opportunities"]
    roi = data["roi_results"]
    roadmap = data["roadmap"]
    knowledge = data["knowledge_assessment"]

    add_cover_page(
        doc,
        company["company_profile"]
    )

    add_executive_summary(
        doc,
        report,
        roi,
        knowledge
    )

    add_company_profile(
        doc,
        company
    )

    add_opportunities(
        doc,
        report
    )

    add_roi(
        doc,
        roi
    )

    add_roadmap(
        doc,
        roadmap
    )

    add_readiness(
        doc,
        knowledge
    )

    add_portfolio(
        doc,
        report,
        roi
    )

    add_recommendations(
        doc,
        report,
        roi,
        knowledge
    )

    buffer = BytesIO()

    doc.save(buffer)

    buffer.seek(0)

    return buffer

def generate_report(data):

    doc = Document()

    company = data["company_data"]
    report = data["opportunities"]
    roi = data["roi_results"]
    roadmap = data["roadmap"]
    knowledge = data["knowledge_assessment"]

    add_cover_page(
        doc,
        company["company_profile"]
    )

    add_executive_summary(
        doc,
        report,
        roi,
        knowledge
    )

    add_company_profile(
        doc,
        company
    )

    add_opportunities(
        doc,
        report
    )

    add_roi(
        doc,
        roi
    )

    add_roadmap(
        doc,
        roadmap
    )

    add_readiness(
        doc,
        knowledge
    )

    add_portfolio(
        doc,
        report,
        roi
    )

    add_recommendations(
        doc,
        report,
        roi,
        knowledge
    )

    buffer = BytesIO()

    doc.save(buffer)

    buffer.seek(0)

    return buffer       